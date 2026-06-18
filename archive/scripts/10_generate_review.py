"""
scripts/10_generate_review.py
------------------------------
Generate a structured systematic review from DB evidence.

Can generate:
  - Species profile review (all evidence for one species)
  - Function review (all species that show 'alarm call' + evidence)
  - Comparative review (zebra finch vs canary — structured comparison)
  - Taxonomic review (all Passeriformes vocal communication)

Output formats:
  - Markdown (.md)
  - Word document (.docx) via python-docx

Usage:
    python scripts/10_generate_review.py --mode species --query "Taeniopygia guttata"
    python scripts/10_generate_review.py --mode function --query "alarm call"
    python scripts/10_generate_review.py --mode compare --query "Taeniopygia guttata,Serinus canaria"
    python scripts/10_generate_review.py --mode taxon --query "Passeriformes"
    python scripts/10_generate_review.py --mode function --query "alarm call" --format docx
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))
from src.db import get_connection

NOW     = datetime.utcnow().strftime("%Y-%m-%d")
OUT_DIR = Path("outputs/reviews")


# ── Data fetchers ─────────────────────────────────────────────────────────────

def get_species_evidence(con, scientific_name: str) -> dict:
    sp = con.execute("""
        SELECT s.*, m.summary, m.image_url, m.conservation_status,
               m.wiki_url
        FROM species s
        LEFT JOIN species_metadata m ON s.species_id = m.species_id
        WHERE s.scientific_name = ?
    """, (scientific_name,)).fetchone()

    if not sp:
        return {}

    sid = sp["species_id"]

    papers = con.execute("""
        SELECT p.paper_id, p.title, p.year, p.doi, p.abstract, p.journal,
               p.open_access, ps.developmental_stage, ps.main_outcome
        FROM papers p
        JOIN paper_species ps ON p.paper_id = ps.paper_id
        WHERE ps.species_id = ?
        ORDER BY p.year DESC NULLS LAST
    """, (sid,)).fetchall()

    claims = con.execute("""
        SELECT claim_type, value,
               COUNT(DISTINCT paper_id) AS paper_count,
               AVG(confidence) AS avg_conf,
               MAX(source) AS best_source
        FROM communication_claims
        WHERE species_id = ?
        GROUP BY claim_type, value
        ORDER BY claim_type, paper_count DESC
    """, (sid,)).fetchall()

    dois = con.execute("""
        SELECT doi, url, title, year FROM open_literature
        WHERE species_id = ? ORDER BY year DESC NULLS LAST
    """, (sid,)).fetchall()

    return {
        "species": dict(sp),
        "papers": [dict(p) for p in papers],
        "claims": [dict(c) for c in claims],
        "open_dois": [dict(d) for d in dois],
    }


def get_function_evidence(con, function_value: str) -> dict:
    """All species + evidence for a given communicative function."""
    rows = con.execute("""
        SELECT s.scientific_name, s.common_name_en, s.class_, s.order_, s.family,
               cc.confidence, cc.paper_id,
               COUNT(DISTINCT cc.paper_id) OVER (PARTITION BY cc.species_id) AS paper_count
        FROM communication_claims cc
        JOIN species s ON cc.species_id = s.species_id
        WHERE cc.claim_type = 'function' AND lower(cc.value) LIKE ?
        ORDER BY paper_count DESC, s.class_, s.scientific_name
    """, (f"%{function_value.lower()}%",)).fetchall()

    papers = con.execute("""
        SELECT DISTINCT p.paper_id, p.title, p.year, p.doi, p.abstract, p.journal,
               s.scientific_name, s.common_name_en, ps.main_outcome
        FROM papers p
        JOIN paper_species ps ON p.paper_id = ps.paper_id
        JOIN species s ON ps.species_id = s.species_id
        JOIN communication_claims cc ON cc.paper_id = p.paper_id AND cc.species_id = s.species_id
        WHERE cc.claim_type = 'function' AND lower(cc.value) LIKE ?
        ORDER BY p.year DESC NULLS LAST
    """, (f"%{function_value.lower()}%",)).fetchall()

    return {
        "function": function_value,
        "species": [dict(r) for r in rows],
        "papers": [dict(p) for p in papers],
    }


def get_comparative_evidence(con, names: list[str]) -> list[dict]:
    return [get_species_evidence(con, n.strip()) for n in names]


def get_taxon_evidence(con, taxon: str) -> dict:
    """All species in a given order or family."""
    species = con.execute("""
        SELECT s.species_id, s.scientific_name, s.common_name_en, s.order_, s.family,
               COUNT(DISTINCT ps.paper_id) AS paper_count,
               COUNT(DISTINCT cc.claim_id) AS claim_count
        FROM species s
        LEFT JOIN paper_species ps ON s.species_id = ps.species_id
        LEFT JOIN communication_claims cc ON s.species_id = cc.species_id
        WHERE s.order_ = ? OR s.family = ? OR s.class_ = ?
        GROUP BY s.species_id
        ORDER BY paper_count DESC, s.scientific_name
    """, (taxon, taxon, taxon)).fetchall()

    return {
        "taxon": taxon,
        "species": [dict(r) for r in species],
    }


# ── Markdown generators ───────────────────────────────────────────────────────

def claims_section(claims: list[dict]) -> str:
    by_type: dict[str, list] = {}
    for c in claims:
        by_type.setdefault(c["claim_type"], []).append(c)

    out = ""
    for ct in ["vocalisation", "context", "function", "method"]:
        items = by_type.get(ct, [])
        if not items:
            continue
        out += f"\n**{ct.capitalize()}s:**\n"
        for item in items:
            pc   = item.get("paper_count", 1)
            conf = item.get("avg_conf", 0.5)
            src  = item.get("best_source", "seed")
            badge = "★" if src == "extraction" else "○"
            out += f"- {badge} {item['value']}  *(n={pc} papers, conf={conf:.2f})*\n"
    return out


def papers_section(papers: list[dict], max_show: int = 20) -> str:
    if not papers:
        return "_No papers in database._\n"
    out = ""
    for i, p in enumerate(papers[:max_show], 1):
        doi_link = f" · [DOI](https://doi.org/{p['doi']})" if p.get("doi") else ""
        oa_badge = " 🔓" if p.get("open_access") else ""
        outcome  = f"\n  > {p['main_outcome']}" if p.get("main_outcome") else ""
        abstract = f"\n  > *{p['abstract'][:200]}…*" if p.get("abstract") else ""
        out += f"{i}. **{p['title']}** ({p.get('year','n.d.')}){oa_badge}{doi_link}\n"
        out += f"   *{p.get('journal','—')}*"
        out += outcome or abstract
        out += "\n\n"
    if len(papers) > max_show:
        out += f"*… and {len(papers)-max_show} more papers in the database.*\n"
    return out


def generate_species_review(data: dict) -> str:
    sp = data["species"]
    papers = data["papers"]
    claims = data["claims"]
    dois   = data["open_dois"]
    n_papers = len(papers)
    n_claims = len(claims)

    md  = f"# {sp['common_name_en'].title()} (*{sp['scientific_name']}*)\n\n"
    md += f"**Systematic review of vocal communication evidence**  \n"
    md += f"Generated: {NOW} · Zoe.Logos-Graph v0.4  \n"
    md += f"Evidence base: {n_papers} papers · {n_claims} communication claims\n\n"
    md += "---\n\n"

    md += "## Taxonomy\n\n"
    md += f"| Field | Value |\n|---|---|\n"
    md += f"| Class | {sp['class_']} |\n"
    md += f"| Order | {sp['order_']} |\n"
    md += f"| Family | {sp['family']} |\n"
    md += f"| Conservation | {sp.get('conservation_status') or '—'} |\n\n"

    if sp.get("summary"):
        md += "## Overview\n\n"
        md += sp["summary"] + "\n\n"
        if sp.get("wiki_url"):
            md += f"[Wikipedia →]({sp['wiki_url']})\n\n"

    md += "## Communication profile\n\n"
    md += "> ★ = extracted from paper · ○ = seed data\n\n"
    md += claims_section(claims) or "_No structured claims available._\n"
    md += "\n"

    md += "## Evidence\n\n"
    md += f"Total papers: **{n_papers}**"
    oa_count = sum(1 for p in papers if p.get("open_access"))
    if oa_count:
        md += f" · Open access: **{oa_count}**"
    md += "\n\n"
    md += papers_section(papers)

    if dois:
        md += "## Open-access literature\n\n"
        for d in dois:
            md += f"- [{d.get('title') or d['doi']}]({d['url']})"
            if d.get("year"):
                md += f" ({d['year']})"
            md += "\n"
        md += "\n"

    md += "---\n*Generated by Zoe.Logos-Graph · Not a substitute for primary literature review*\n"
    return md


def generate_function_review(data: dict) -> str:
    fn = data["function"]
    species = data["species"]
    papers  = data["papers"]

    # Deduplicate species
    seen_sci = set()
    unique_sps = []
    for s in species:
        if s["scientific_name"] not in seen_sci:
            seen_sci.add(s["scientific_name"])
            unique_sps.append(s)

    md  = f"# {fn.title()} — Cross-species Evidence Review\n\n"
    md += f"**Systematic review of '{fn}' in animal vocal communication**  \n"
    md += f"Generated: {NOW} · Zoe.Logos-Graph v0.4  \n"
    md += f"Evidence base: {len(unique_sps)} species · {len(papers)} papers\n\n"
    md += "---\n\n"

    md += f"## Species showing '{fn}'\n\n"
    md += "| Species | Common name | Class | Order | Papers |\n|---|---|---|---|---|\n"
    for s in unique_sps:
        md += f"| *{s['scientific_name']}* | {s['common_name_en']} | {s['class_']} | {s['order_']} | {s.get('paper_count',0)} |\n"
    md += "\n"

    # Taxonomic distribution
    by_class: dict[str, list] = {}
    for s in unique_sps:
        by_class.setdefault(s["class_"], []).append(s["common_name_en"])
    md += "## Taxonomic distribution\n\n"
    for cls, sps in sorted(by_class.items()):
        md += f"- **{cls}**: {', '.join(sps)}\n"
    md += "\n"

    md += "## Papers\n\n"
    md += papers_section(papers)

    md += "---\n*Generated by Zoe.Logos-Graph · Not a substitute for primary literature review*\n"
    return md


def generate_comparative_review(datasets: list[dict]) -> str:
    names = [d["species"]["scientific_name"] for d in datasets if d]
    md  = f"# Comparative Review: {' vs '.join(names)}\n\n"
    md += f"Generated: {NOW} · Zoe.Logos-Graph v0.4\n\n---\n\n"

    # Side-by-side claims table
    claim_types = ["vocalisation","context","function","method"]
    for ct in claim_types:
        md += f"### {ct.capitalize()}s\n\n"
        md += "| " + " | ".join(d["species"]["common_name_en"] for d in datasets if d) + " |\n"
        md += "|" + "---|" * len([d for d in datasets if d]) + "\n"
        # Collect all values
        all_vals = set()
        sp_vals: dict[str, set] = {}
        for d in datasets:
            if not d:
                continue
            sci = d["species"]["scientific_name"]
            sp_vals[sci] = set(c["value"] for c in d["claims"] if c["claim_type"] == ct)
            all_vals |= sp_vals[sci]
        for val in sorted(all_vals):
            row = "| "
            for d in datasets:
                if not d:
                    continue
                sci = d["species"]["scientific_name"]
                row += ("✓" if val in sp_vals.get(sci, set()) else "—") + " | "
            md += row + f"*{val}*\n"
        md += "\n"

    md += "## Individual profiles\n\n"
    for d in datasets:
        if not d:
            continue
        sp = d["species"]
        md += f"### *{sp['scientific_name']}* ({sp['common_name_en']})\n\n"
        md += f"Papers: {len(d['papers'])} · Claims: {len(d['claims'])}\n\n"
        md += claims_section(d["claims"]) + "\n"

    md += "---\n*Generated by Zoe.Logos-Graph · Not a substitute for primary literature review*\n"
    return md


def generate_taxon_review(data: dict) -> str:
    taxon = data["taxon"]
    species = data["species"]

    md  = f"# {taxon} — Vocal Communication Review\n\n"
    md += f"Generated: {NOW} · Zoe.Logos-Graph v0.4  \n"
    md += f"Species covered: {len(species)}\n\n---\n\n"

    md += "## Species in database\n\n"
    md += "| Species | Common name | Family | Papers | Claims |\n|---|---|---|---|---|\n"
    for s in species:
        md += f"| *{s['scientific_name']}* | {s['common_name_en']} | {s['family']} | {s['paper_count']} | {s['claim_count']} |\n"
    md += "\n"

    total_papers = sum(s["paper_count"] for s in species)
    total_claims = sum(s["claim_count"] for s in species)
    md += f"**Total**: {len(species)} species · {total_papers} papers · {total_claims} claims\n\n"
    md += "---\n*Generated by Zoe.Logos-Graph · Not a substitute for primary literature review*\n"
    return md


# ── Export ────────────────────────────────────────────────────────────────────

def export_docx(markdown_text: str, output_path: Path) -> bool:
    """Convert markdown to docx using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  python-docx not installed: pip install python-docx")
        return False

    doc = Document()

    # Style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in markdown_text.splitlines():
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Pt(24)
        elif line.startswith("---"):
            doc.add_paragraph("─" * 60)
        elif line.startswith("|"):
            # Skip tables for now
            doc.add_paragraph(line)
        elif line.strip():
            doc.add_paragraph(line)

    doc.save(str(output_path))
    return True


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate systematic review from DB evidence.")
    parser.add_argument("--mode",   required=True, choices=["species","function","compare","taxon"])
    parser.add_argument("--query",  required=True, help="Species name, function, or taxon")
    parser.add_argument("--format", default="md",  choices=["md","docx","both"])
    args = parser.parse_args()

    con = get_connection()
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    print(f"\nGenerating {args.mode} review for: {args.query}\n")

    if args.mode == "species":
        data = get_species_evidence(con, args.query)
        if not data:
            print(f"Species not found: {args.query}")
            sys.exit(1)
        md_text = generate_species_review(data)
        slug = args.query.replace(" ","_").lower()

    elif args.mode == "function":
        data = get_function_evidence(con, args.query)
        md_text = generate_function_review(data)
        slug = f"function_{args.query.replace(' ','_').lower()}"

    elif args.mode == "compare":
        names = args.query.split(",")
        data  = get_comparative_evidence(con, names)
        md_text = generate_comparative_review(data)
        slug = "compare_" + "_vs_".join(n.strip().split()[0].lower() for n in names)

    elif args.mode == "taxon":
        data = get_taxon_evidence(con, args.query)
        md_text = generate_taxon_review(data)
        slug = f"taxon_{args.query.replace(' ','_').lower()}"

    # Save
    md_path = OUT_DIR / f"{slug}_{NOW}.md"
    md_path.write_text(md_text, encoding="utf-8")
    print(f"✓ Markdown: {md_path}")

    if args.format in ("docx","both"):
        docx_path = OUT_DIR / f"{slug}_{NOW}.docx"
        if export_docx(md_text, docx_path):
            print(f"✓ Word doc: {docx_path}")

    print(f"\nPreview (first 400 chars):\n")
    print(md_text[:400])
    print("…\n")


if __name__ == "__main__":
    main()
